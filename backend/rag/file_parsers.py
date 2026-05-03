"""
File parsers for multiple formats: .docx, .pdf, .txt
"""

import logging
import re
from pathlib import Path
from typing import Tuple, Optional

logger = logging.getLogger(__name__)


class FileParserBase:
    """Base class for file parsers"""
    
    def __init__(self):
        self.supported_formats = []
    
    def extract_text(self, file_path: str) -> str:
        """Extract text from file"""
        raise NotImplementedError


class DocxParser(FileParserBase):
    """Parse .docx files"""
    
    def __init__(self):
        super().__init__()
        self.supported_formats = ['.docx']
        try:
            from docx import Document
            self.Document = Document
        except ImportError:
            raise ImportError("python-docx not installed. Run: pip install python-docx")
    
    def extract_text(self, file_path: str) -> str:
        """
        Extract text from DOCX file
        
        Args:
            file_path: Path to .docx file
            
        Returns:
            Extracted and cleaned text
        """
        try:
            doc = self.Document(file_path)
            text_parts = []
            
            for para in doc.paragraphs:
                if para.text.strip():
                    text_parts.append(para.text)
            
            # Also extract from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text.strip():
                            text_parts.append(cell.text)
            
            full_text = "\n\n".join(text_parts)
            return self._clean_text(full_text)
            
        except Exception as e:
            logger.error(f"Error parsing DOCX {file_path}: {str(e)}")
            raise
    
    @staticmethod
    def _clean_text(text: str) -> str:
        """Clean extracted text"""
        # Remove extra spaces
        text = re.sub(r' +', ' ', text)
        # Normalize line breaks
        text = re.sub(r'\n\n+', '\n\n', text)
        # Strip leading/trailing whitespace
        text = text.strip()
        return text


class PdfParser(FileParserBase):
    """Parse .pdf files"""
    
    def __init__(self, use_pdfplumber: bool = True):
        super().__init__()
        self.supported_formats = ['.pdf']
        
        if use_pdfplumber:
            try:
                import pdfplumber
                self.pdf_lib = pdfplumber
                self.lib_name = 'pdfplumber'
            except ImportError:
                logger.warning("pdfplumber not installed, trying PyPDF2...")
                try:
                    import PyPDF2
                    self.pdf_lib = PyPDF2
                    self.lib_name = 'PyPDF2'
                except ImportError:
                    raise ImportError(
                        "Neither pdfplumber nor PyPDF2 installed. "
                        "Run: pip install pdfplumber or pip install PyPDF2"
                    )
        else:
            try:
                import PyPDF2
                self.pdf_lib = PyPDF2
                self.lib_name = 'PyPDF2'
            except ImportError:
                raise ImportError("PyPDF2 not installed. Run: pip install PyPDF2")
    
    def extract_text(self, file_path: str) -> str:
        """
        Extract text from PDF file
        
        Args:
            file_path: Path to .pdf file
            
        Returns:
            Extracted and cleaned text
        """
        try:
            if self.lib_name == 'pdfplumber':
                return self._extract_pdfplumber(file_path)
            else:
                return self._extract_pypdf2(file_path)
                
        except Exception as e:
            logger.error(f"Error parsing PDF {file_path}: {str(e)}")
            raise
    
    def _extract_pdfplumber(self, file_path: str) -> str:
        """Extract using pdfplumber"""
        text_parts = []
        
        with self.pdf_lib.open(file_path) as pdf:
            for page in pdf.pages:
                page_text = page.extract_text()
                if page_text:
                    text_parts.append(page_text)
        
        full_text = "\n\n".join(text_parts)
        return self._clean_text(full_text)
    
    def _extract_pypdf2(self, file_path: str) -> str:
        """Extract using PyPDF2"""
        text_parts = []
        
        with open(file_path, 'rb') as pdf_file:
            reader = self.pdf_lib.PdfReader(pdf_file)
            for page in reader.pages:
                page_text = page.extract_text()
                if page_text:
                    text_parts.append(page_text)
        
        full_text = "\n\n".join(text_parts)
        return self._clean_text(full_text)
    
    @staticmethod
    def _clean_text(text: str) -> str:
        """Clean extracted text"""
        # Remove extra spaces
        text = re.sub(r' +', ' ', text)
        # Normalize line breaks
        text = re.sub(r'\n\n+', '\n\n', text)
        # Remove form feed characters
        text = text.replace('\x0c', '')
        # Strip leading/trailing whitespace
        text = text.strip()
        return text


class TxtParser(FileParserBase):
    """Parse .txt files"""
    
    def __init__(self):
        super().__init__()
        self.supported_formats = ['.txt']
    
    def extract_text(self, file_path: str) -> str:
        """
        Extract text from TXT file
        
        Args:
            file_path: Path to .txt file
            
        Returns:
            Extracted and cleaned text
        """
        try:
            with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
                text = f.read()
            
            return self._clean_text(text)
            
        except Exception as e:
            logger.error(f"Error parsing TXT {file_path}: {str(e)}")
            raise
    
    @staticmethod
    def _clean_text(text: str) -> str:
        """Clean extracted text"""
        # Remove extra spaces
        text = re.sub(r' +', ' ', text)
        # Normalize line breaks
        text = re.sub(r'\n\n+', '\n\n', text)
        # Strip leading/trailing whitespace
        text = text.strip()
        return text


class UniversalFileParser:
    """Universal parser supporting multiple file formats"""
    
    def __init__(self):
        """Initialize all available parsers"""
        self.parsers = {}
        
        # DOCX parser
        try:
            parser = DocxParser()
            for fmt in parser.supported_formats:
                self.parsers[fmt] = parser
            logger.info("✓ DOCX parser available")
        except ImportError as e:
            logger.warning(f"DOCX parser unavailable: {str(e)}")
        
        # PDF parser
        try:
            parser = PdfParser(use_pdfplumber=True)
            for fmt in parser.supported_formats:
                self.parsers[fmt] = parser
            logger.info(f"✓ PDF parser available ({parser.lib_name})")
        except ImportError as e:
            logger.warning(f"PDF parser unavailable: {str(e)}")
        
        # TXT parser (always available)
        parser = TxtParser()
        for fmt in parser.supported_formats:
            self.parsers[fmt] = parser
        logger.info("✓ TXT parser available")
    
    def get_supported_formats(self) -> list:
        """Get list of supported file formats"""
        return list(set(fmt for fmt in self.parsers.keys()))
    
    def is_supported(self, file_path: str) -> bool:
        """Check if file format is supported"""
        file_ext = Path(file_path).suffix.lower()
        return file_ext in self.parsers
    
    def extract_text(self, file_path: str) -> str:
        """
        Extract text from file
        
        Args:
            file_path: Path to file
            
        Returns:
            Extracted text
            
        Raises:
            ValueError: If file format not supported
        """
        file_path = str(file_path)
        file_ext = Path(file_path).suffix.lower()
        
        if file_ext not in self.parsers:
            supported = ", ".join(self.get_supported_formats())
            raise ValueError(
                f"File format {file_ext} not supported. "
                f"Supported formats: {supported}"
            )
        
        parser = self.parsers[file_ext]
        logger.info(f"Parsing {Path(file_path).name} using {parser.__class__.__name__}")
        
        return parser.extract_text(file_path)
    
    def parse_file(self, file_path: str) -> Tuple[str, str]:
        """
        Parse file and return (file_name, extracted_text)
        
        Args:
            file_path: Path to file
            
        Returns:
            Tuple of (file_name, text)
        """
        file_path = Path(file_path)
        text = self.extract_text(str(file_path))
        return (file_path.name, text)


# Module-level parser instance
_parser_instance = None


def get_universal_parser() -> UniversalFileParser:
    """Get or create universal parser instance"""
    global _parser_instance
    if _parser_instance is None:
        _parser_instance = UniversalFileParser()
    return _parser_instance


def extract_text(file_path: str) -> str:
    """
    Convenience function to extract text from any supported file
    
    Args:
        file_path: Path to file
        
    Returns:
        Extracted text
    """
    parser = get_universal_parser()
    return parser.extract_text(file_path)
